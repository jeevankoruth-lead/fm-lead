import os
from docx import Document
from docx.shared import Inches
from glob import glob

SATIRE_ROOT = os.path.join(os.path.dirname(__file__), '..', 'content', 'muser', 'satire')

# Helper to find images in a satire post folder
def get_images(post_dir):
    images_dir = os.path.join(post_dir, 'images')
    if not os.path.isdir(images_dir):
        return []
    # Accept common image formats
    return sorted(glob(os.path.join(images_dir, '*.[pjgPJG]*')))  # crude, but works for png/jpg/jpeg/gif

# Helper to reduce bullet points and consolidate text
def consolidate_docx(docx_path):
    doc = Document(docx_path)
    new_doc = Document()
    bullet_buffer = []
    for para in doc.paragraphs:
        if para.style.name.startswith('List'):
            bullet_buffer.append(para.text)
        else:
            if bullet_buffer:
                # Consolidate bullets into a single paragraph
                new_doc.add_paragraph(' • '.join(bullet_buffer))
                bullet_buffer = []
            new_doc.add_paragraph(para.text)
    if bullet_buffer:
        new_doc.add_paragraph(' • '.join(bullet_buffer))
    # Insert images at the end
    post_dir = os.path.dirname(docx_path)
    images = get_images(post_dir)
    for img_path in images:
        try:
            new_doc.add_picture(img_path, width=Inches(4.5))
        except Exception as e:
            print(f'Could not add image {img_path}: {e}')
    # Overwrite the original docx
    new_doc.save(docx_path)
    print(f'Updated: {docx_path}')

def main():
    satire_docs = glob(os.path.join(SATIRE_ROOT, '*', 'index.docx'))
    for docx_path in satire_docs:
        consolidate_docx(docx_path)

if __name__ == '__main__':
    main()
