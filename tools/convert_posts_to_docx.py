import os
import glob
from markdown import markdown
from docx import Document
from bs4 import BeautifulSoup

# Root content directory (one level up from tools)
CONTENT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'content'))

# Find all index.md files in post folders
def find_post_markdowns():
    return glob.glob(os.path.join(CONTENT_ROOT, '**', 'index.md'), recursive=True)

def md_to_docx(md_path):
    # Read markdown
    with open(md_path, encoding='utf-8') as f:
        md_text = f.read()
    # Convert markdown to HTML
    html = markdown(md_text)
    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')
    # Create Word document
    doc = Document()
    for elem in soup.descendants:
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text())
        elif elem.name == 'ul':
            for li in elem.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif elem.name == 'ol':
            for li in elem.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
    # Save .docx in same folder
    docx_path = os.path.join(os.path.dirname(md_path), 'index.docx')
    doc.save(docx_path)
    print(f'Converted: {md_path} -> {docx_path}')


def main():
    md_files = find_post_markdowns()
    if not md_files:
        print(f'No Markdown files found in {CONTENT_ROOT}')
    for md_path in md_files:
        try:
            md_to_docx(md_path)
        except Exception as e:
            print(f'Error converting {md_path}: {e}')

if __name__ == '__main__':
    main()
